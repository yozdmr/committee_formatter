from docx import Document

def committee_type(paragraph, committee_index):
    if paragraph.runs[0].bold and paragraph.runs[0].italic:
        if 'Ad Hoc' in paragraph.text:
            committee_index = 1
        return True
    return False

def committee(paragraph, c):
    if paragraph.runs[0].bold and paragraph.runs[0].underline:
        name = paragraph.text
        name = " ".join(name.split())
        c['name'] = name
        return True
    return False

def category(paragraph, paragraphs, i, c):
    if paragraph.runs[0].bold and not paragraph.runs[0].underline:
        text = paragraph.text
        if 'Note' in text:
            c['note'] = paragraph.text
        elif 'Chair' in text:
            # Account for co-chairs
            chair = paragraph.text
            chair = " ".join(chair.split())
            if chair == 'Chair and':
                chair = 'LIAISON'
            c['chair'] = chair
        elif 'Members' in text:
            # Add the following lines to c['members'] until 
            # paragraph.runs[0].bold is true again
            c['members'] = []
            index = i

            while index+1 < len(paragraphs):
                member = paragraphs[index].text
                member = " ".join(member.split())
                c['members'].append(member)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
        elif 'Liaison' in text:
            c['liaison'] = paragraph.text
        elif 'Alternates' in text:
            # Add the following lines to c['alternates'] until 
            # paragraph.runs[0].bold is true again
            c['alternates'] = []
            index = i
            while index+1 < len(paragraphs):
                c['alternates'].append(paragraphs[index].text)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
        elif 'Committee Charge' in text:
            # Add the following lines to c['charge'] until 
            # paragraph.runs[0].bold is true again
            c['charge'] = []
            index = i
            while index+1 < len(paragraphs):
                c['charge'].append(paragraphs[index].text)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
            return True
    return False


def run_doc(file_path):
    doc = Document(file_path)
    committee_list = [
    [{}],
    [{}],
]
    committee_index = 0
    c_index = 0

    for index, paragraph in enumerate(doc.paragraphs):
        curr_c = committee_list[committee_index][c_index]
        text = paragraph.text
        runs = len(paragraph.runs)

        if runs > 0:
            # Here reference above functions
            a = committee_type(paragraph, committee_index)
            b = committee(paragraph, curr_c)
            end_c = category(paragraph, doc.paragraphs, index, curr_c)
            if end_c:
                committee_list[committee_index].append({})
                c_index+=1
    
    return committee_list

committee_list = run_doc("c.docx")

import json
with open('committee_list.json', 'w') as f:
    json.dump(committee_list, f, indent=4)


''' committee dict format
# Example: The Academic Appeals Committee
c = {
    'name': "Academic Appeals Committee",
    'note': None
    'chair': "EJ Ume (ECO) (2025)",
    'members': [
        'Geoff Zoeckler (ESP) (2026)'
	    'David Gempesaw (FIN) (2025)'
        'Michael Gowins (ISA) (2026)'
	    'Peter Nguyen (MKT) (2025)'	 
	    'Bill Moser (ACC) (2026)'
    ],
    'liaison': 'Chanelle White',
    'alternates': [
        'Sydney Shu (ACC) (2026)'
        'Nam Vu (ECO) (2025)'
        'Chris Sutter (ESP) (2026)'
	    'Tyler Henry (FIN) (2025)'
	    'Vivian Chen (ISA) (2026)'
        'Darryl Rice (MGT) (2025)'
        'Eric Stenstrom (MKT) (2025)'
    ],
    'charge': [
        'Consider academic grievances as outlined in the Student Handbook.'
        'The Committee will meet to address its annual charge. Meeting minutes will be taken.'
        'The Committee Chair will provide a summative year-end report of Committee activities to the Associate Dean for Curriculum. This report will be added as a consent item to the agenda of the first faculty meeting of the following academic year.'
    ]
}
'''