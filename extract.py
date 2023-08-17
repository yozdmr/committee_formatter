from docx import Document
import json
import os

def committee(paragraph, c):
    if paragraph.runs[0].bold and paragraph.runs[0].underline:
        name = paragraph.text
        name = " ".join(name.split())
        c['name'] = name
        return True
    return False

def category(paragraph, paragraphs, i, c):
    if paragraph.runs[0].bold and not paragraph.runs[0].underline:
        text = paragraph.text
        if 'Note:' in text:
            c['note'] = paragraph.text
        elif 'Chair:' in text or 'Co-Chairs:' in text or "Chair and" in text:
            c['chairs'] = []
            index = i

            while index+1 < len(paragraphs):
                chair = paragraphs[index].text
                chair = " ".join(chair.split())
                c['chairs'].append(chair)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
        elif 'Members:' in text:
            c['members'] = []
            index = i

            while index+1 < len(paragraphs):
                member = paragraphs[index].text
                member = " ".join(member.split())
                c['members'].append(member)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
        elif 'Liaison:' in text:
            liaison = " ".join(text.split())
            c['liaison'] = liaison
        elif 'Alternates:' in text:
            c['alternates'] = []
            index = i
            while index+1 < len(paragraphs):
                alt = paragraphs[index].text
                alt = " ".join(alt.split())
                c['alternates'].append(alt)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
        elif 'Committee Charge:' in text:
            c['charge'] = []
            index = i
            while index+1 < len(paragraphs):
                c['charge'].append(paragraphs[index].text)
                if len(paragraphs[index+1].runs) > 0 and paragraphs[index+1].runs[0].bold:
                    break
                index += 1
            return True
    return False

def run_doc(file_name):
    file_path = os.path.join(os.getcwd(), "input", file_name)
    doc = Document(file_path)
    committee_list = [
        [{}],
        [{}],
    ]
    committee_index = 0
    c_index = 0

    for index, paragraph in enumerate(doc.paragraphs):
        curr_c = committee_list[committee_index][c_index]
        text = paragraph.text
        runs = len(paragraph.runs)

        if runs > 0:
            if paragraph.runs[0].bold and paragraph.runs[0].italic:
                if 'Ad Hoc' in paragraph.text:
                    committee_index = 1
                    c_index = 0
            committee(paragraph, curr_c)
            end_c = category(paragraph, doc.paragraphs, index, curr_c)
            if end_c:
                committee_list[committee_index].append({})
                c_index+=1

    file_name = file_name.split(".")[0]
    with open(f'bin/{file_name}.json', 'w') as f:
        json.dump(committee_list, f, indent=4)
    
    return committee_list